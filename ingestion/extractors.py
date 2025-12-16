# ingestion/extractors.py
import os
import fitz  # PyMuPDF
import csv
from typing import List, Optional
import logging
from models.data_models import Document, SmartChunk

logger = logging.getLogger("DocumentExtractors")

class DocumentExtractor:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.txt': self._extract_text,
            '.csv': self._extract_csv,
            '.docx': self._extract_docx,
            '.xlsx': self._extract_excel,
        }
    
    def extract(self, document_path: str) -> Optional[Document]:
        file_ext = os.path.splitext(document_path)[1].lower()
        
        if file_ext not in self.supported_formats:
            logger.warning(f"Unsupported file format: {file_ext}")
            return None
        
        try:
            file_size = os.path.getsize(document_path)
            last_modified = os.path.getmtime(document_path)
            
            document = Document(
                path=document_path,
                document_id=os.path.basename(document_path),
                file_type=file_ext,
                size=file_size,
                last_modified=last_modified
            )
            
            # Extraer texto usando el método específico
            extractor_func = self.supported_formats[file_ext]
            text_content = extractor_func(document_path)
            
            if text_content:
                document.chunks = self._chunk_text(text_content, document_path)
            
            return document
            
        except Exception as e:
            logger.error(f"Error extracting {document_path}: {e}")
            return None
    
    def _extract_pdf(self, file_path: str) -> str:
        text = ""
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.get_text()
        return text
    
    def _extract_text(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _extract_csv(self, file_path: str) -> str:
        lines = []
        with open(file_path, 'r', encoding='utf-8', newline='') as f:
            reader = csv.reader(f)
            for row in reader:
                lines.append("\t".join([str(x) for x in row]))
        return "\n".join(lines)
    
    def _extract_docx(self, file_path: str) -> str:
        # Placeholder para extracción de DOCX
        # Necesitarías instalar python-docx
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except ImportError:
            logger.error("python-docx not installed for DOCX extraction")
            return ""
    
    def _extract_excel(self, file_path: str) -> str:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"[Hoja] {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    parts.append("\t".join(["" if v is None else str(v) for v in row]))
            return "\n".join(parts)
        except Exception as e:
            logger.error(f"Excel parse error: {e}")
            return ""
    
    def _chunk_text(self, text: str, source_path: str) -> List[SmartChunk]:
        # Implementación simple de chunking
        # En una implementación real, usarías un splitter más inteligente
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), 200):  # Chunks de ~200 palabras
            chunk_text = " ".join(words[i:i+200])
            chunk = SmartChunk(
                text=chunk_text,
                chunk_id=f"{source_path}_chunk_{i}",
                source_document=source_path,
                document_id=os.path.basename(source_path),
                metadata={
                    "chunk_index": i,
                    "word_count": len(chunk_text.split())
                }
            )
            chunks.append(chunk)
        
        return chunks
